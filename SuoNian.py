import sys
import os

def install_package(package):
    print(f"Installing {package}...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx not found. Installing...")
    install_package("python-docx")
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import re

def markdown_to_docx(md_file, docx_file):
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found.")
        return

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Set default font to support Chinese if possible, though docx handles it reasonably well usually.
    # We'll just rely on defaults for now to keep it simple and robust.
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        else:
            # Basic formatting processing
            # Bold **text**
            p = doc.add_paragraph()
            
            # Simple bold parsing: split by **
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base_dir, "final_report.md")
    docx_path = os.path.join(base_dir, "final_report.docx")
    markdown_to_docx(md_path, docx_path)
